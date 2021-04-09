import requests, re
from math import *
from bs4 import BeautifulSoup    # HTML을 파싱하는 모듈
import os, random, itertools, time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

### 시간 측정 ### 
first_time = time.time() 

## 파일 생성##
templit = Document("English_Composition.docx")
check = Document()

def googling(key1=str):
    url = "http://www.google.com/search?q=" + key1
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:57.0) Gecko/20100101 Firefox/57.0'}
    
    html = requests.get(url, headers=headers ).text
    soup = BeautifulSoup(html, 'html.parser')    
    div = soup.find('div', id='result-stats')

    if not div:
        return 0 
    result_characters = " ".join(div.text.replace(",","")).split(" ")

    for order in range(len(result_characters)):
        word = result_characters[order]
        if not word in [str(number) for number in range(0,10)]:
            if word == "개":
                result_characters = result_characters[:order]
                break
            else:
                result_characters[order] = ""

    result = str(round(log(int("".join([word for word in result_characters if word])),10),3))
    return result

def creatification(tuple_words, operator = ""):
    string = ""
    for word in tuple_words: 
        string +=  " " + word + operator #operator.pop(0) 을 하면 다음 함수 operator 리스트가 줄어있다.
    if operator:
        string = string[:-1]
    else:
        string = "'" + string.rstrip().lstrip() + "'"
    print(string)
    return string

### 워드 읽기 ###

def reading(templit):

    basket = []
    for row, paragraph in enumerate(templit.paragraphs):
        sentences = [sentence.rstrip().lstrip() for sentence in paragraph.text.replace('\n',' ').split(".")]
        for sentence in sentences:
            while "  " in sentence:
                sentence.replace("  ", " ")
            parts = sentence.split(" ")
            for i in range(len(parts)): 
                if len(parts) == i+1 : #2단계 분석
                    pass
                else :
                    friends = (parts[i],parts[i+1])
                    basket.append(friends)
                if len(parts) <= i+2 : #3단계 분석 + 한 단어 건너뛰기 
                    pass
                else : 
                    friends = (parts[i],parts[i+1],parts[i+2])
                    friends2 = (parts[i],parts[i+2],"*") 
                    basket.append(friends)
                    basket.append(friends2)
    basket.sort()
    return basket

def interpret(basket):
    new_basket = []
    for information in basket:
        if len(information) == 3:
            if information[2] == "*":
                information = information + (googling(creatification(information[:2],"*")),)
            else :
                information = information + (googling(creatification(information)),)
        elif len(information) == 2:
            information = information + (googling(creatification(information)),)
        new_basket.append(information)
    return new_basket

def righting(check,new_basket):

    table = check.add_table(rows = 4,cols = len(new_basket))
    table.style = check.styles['Table Grid']
    for i in range(len(new_basket)):
        LR_chocolate = table.columns[i].cells
        for j in range(len(new_basket[i])):
            LR_chocolate[j].paragraphs[0].add_run(new_basket[i][j])
            LR_chocolate[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    print(new_basket)
    check.save("check.docx")
    return table

basket = reading(templit)
new_basket = interpret(basket)
table = righting(check,new_basket)


# Time Check 
passed_time = time.time() - first_time
print(passed_time)



